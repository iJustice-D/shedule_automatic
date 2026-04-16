from __future__ import annotations

from pathlib import Path

from docx import Document
from sqlmodel import Session

from app.core.timetable import DAYS, PAIR_NUMBERS
from app.services.exporters.context import build_schedule_context


class DocxExporter:
    def export(self, session: Session, schedule_id: int, output_path: Path) -> Path:
        context = build_schedule_context(session, schedule_id)
        document = Document()
        for title, grid in context["group_grids"].items():
            document.add_heading(title, level=2)
            document.add_paragraph("Основное расписание")
            table = document.add_table(rows=1 + len(DAYS), cols=1 + len(PAIR_NUMBERS))
            table.cell(0, 0).text = "День / Пара"
            for column, pair_number in enumerate(PAIR_NUMBERS, start=1):
                table.cell(0, column).text = context["pair_headers"][pair_number]
            for row, (day_of_week, _day_key) in enumerate(DAYS.items(), start=1):
                table.cell(row, 0).text = context["day_labels"][day_of_week]
                for column, pair_number in enumerate(PAIR_NUMBERS, start=1):
                    table.cell(row, column).text = grid.get((day_of_week, pair_number), "")
            document.add_paragraph("Дополнительные онлайн-занятия")
            online_rows = context["group_online_rows"].get(title, [])
            online_table = document.add_table(rows=1 + max(len(online_rows), 1), cols=6)
            for column, label in enumerate(["День", "Онлайн-слот", "Предмет", "Преподаватель", "Формат", "Недели"]):
                online_table.cell(0, column).text = label
            if online_rows:
                for row, item in enumerate(online_rows, start=1):
                    online_table.cell(row, 0).text = item["day"]
                    online_table.cell(row, 1).text = item["online_slot"]
                    online_table.cell(row, 2).text = item["subject"]
                    online_table.cell(row, 3).text = item["teacher"]
                    online_table.cell(row, 4).text = item["format"]
                    online_table.cell(row, 5).text = item["weeks"]
            else:
                online_table.cell(1, 0).text = "Онлайн-занятия не назначены"
            document.add_paragraph("")
        document.add_heading("Баланс нагрузки преподавателей", level=2)
        balance_rows = context.get("teacher_balance_rows", [])
        balance_table = document.add_table(rows=1 + max(len(balance_rows), 1), cols=5)
        for column, label in enumerate(["Преподаватель", "Семестр 3", "Семестр 4", "Отклонение", "На уточнении"]):
            balance_table.cell(0, column).text = label
        if balance_rows:
            for row, item in enumerate(balance_rows, start=1):
                balance_table.cell(row, 0).text = str(item.get("teacher_name", ""))
                balance_table.cell(row, 1).text = str(item.get("semester_3_pairs", ""))
                balance_table.cell(row, 2).text = str(item.get("semester_4_pairs", ""))
                balance_table.cell(row, 3).text = str(item.get("normalized_balance_score", ""))
                balance_table.cell(row, 4).text = str(item.get("pending_rows", ""))
        else:
            balance_table.cell(1, 0).text = "Данные отсутствуют"
        document.add_paragraph("")
        document.add_heading("Вакансии и неразрешённые строки", level=2)
        unresolved_rows = context.get("unresolved_weekly_rows_report", [])
        unresolved_table = document.add_table(rows=1 + max(len(unresolved_rows), 1), cols=6)
        for column, label in enumerate(["Группа", "Семестр", "Предмет", "Подгруппа", "Состояние", "Преподаватели"]):
            unresolved_table.cell(0, column).text = label
        if unresolved_rows:
            for row, item in enumerate(unresolved_rows, start=1):
                unresolved_table.cell(row, 0).text = str(item.get("group", ""))
                unresolved_table.cell(row, 1).text = str(item.get("semester", ""))
                unresolved_table.cell(row, 2).text = str(item.get("subject", ""))
                unresolved_table.cell(row, 3).text = str(item.get("subgroup", ""))
                unresolved_table.cell(row, 4).text = str(item.get("assignment_state", ""))
                unresolved_table.cell(row, 5).text = str(item.get("teacher_names", ""))
        else:
            unresolved_table.cell(1, 0).text = "Неразрешённых строк нет"
        document.save(output_path)
        return output_path
